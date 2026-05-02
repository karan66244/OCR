from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
from img2table.document import Image
from img2table.ocr import TesseractOCR
import pytesseract
from core.config import settings
from openpyxl import Workbook
import os


def export_to_docx(text: str, filename: str) -> BytesIO:
    """
    Export extracted text to DOCX format with formatting.
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading('Extracted Text', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    metadata = doc.add_paragraph()
    metadata.add_run(f"Filename: {filename}\n").italic = True
    metadata_format = metadata.paragraph_format
    metadata_format.space_before = Pt(6)
    metadata_format.space_after = Pt(12)
    
    # Add separator
    doc.add_paragraph('_' * 80)
    
    # Parse text and add paragraphs with proper formatting
    paragraphs = text.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            para = doc.add_paragraph(para_text.strip())
            para_format = para.paragraph_format
            para_format.line_spacing = 1.5
            para_format.space_after = Pt(12)
            para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Save to BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def export_to_pdf(text: str, filename: str) -> BytesIO:
    """
    Export extracted text to PDF format.
    """
    output = BytesIO()
    doc = SimpleDocTemplate(
        output,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    
    # Create story (content)
    story = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=RGBColor(0, 51, 102),
        spaceAfter=30,
        alignment=TA_LEFT,
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['BodyText'],
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=12,
    )
    
    # Add title
    story.append(Paragraph("Extracted Text", title_style))
    story.append(Spacer(1, 12))
    
    # Add metadata
    metadata_text = f"<i>Filename: {filename}</i>"
    story.append(Paragraph(metadata_text, styles['Italic']))
    story.append(Spacer(1, 12))
    
    # Add separator
    story.append(Paragraph("_" * 80, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Parse and add text
    paragraphs = text.split('\n\n')
    for idx, para_text in enumerate(paragraphs):
        if para_text.strip():
            # Replace newlines within paragraph
            formatted_text = para_text.strip().replace('\n', '<br/>')
            story.append(Paragraph(formatted_text, body_style))
            
            # Add page break after every 5 paragraphs to maintain readability
            if (idx + 1) % 5 == 0:
                story.append(PageBreak())
    
    # Build PDF
    doc.build(story)
    output.seek(0)
    return output


def export_to_txt(text: str) -> BytesIO:
    """
    Export extracted text to TXT format.
    """
    output = BytesIO()
    output.write(text.encode('utf-8'))
    output.seek(0)
    return output


# Configure tesseract path for table extraction
if settings.TESSERACT_PATH:
    pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_PATH
    tesseract_dir = os.path.dirname(settings.TESSERACT_PATH)
    if tesseract_dir:
        os.environ["PATH"] = f"{tesseract_dir}{os.pathsep}{os.environ.get('PATH', '')}"
        tessdata_dir = os.path.join(tesseract_dir, "tessdata")
        if os.path.isdir(tessdata_dir):
            os.environ.setdefault("TESSDATA_PREFIX", tessdata_dir)


def export_tables_to_xlsx(image_bytes: bytes, filename: str) -> BytesIO:
    """
    Export detected tables from an image to XLSX (one sheet per table).
    """
    ocr = TesseractOCR(n_threads=1, lang="eng")
    image = Image(src=image_bytes)

    tables = image.extract_tables(
        ocr=ocr,
        implicit_rows=True,
        min_confidence=50,
    )

    if not tables:
        raise ValueError("No tables detected in the image")

    workbook = Workbook()
    workbook.remove(workbook.active)

    for idx, table in enumerate(tables, start=1):
        sheet = workbook.create_sheet(title=f"Table {idx}")
        
        try:
            df = table.df
            
            # Set header from the first row of the DataFrame
            header = df.iloc[0]
            df = df[1:]
            df.columns = header
            
            # Write header
            for col_idx, col_name in enumerate(header, 1):
                sheet.cell(row=1, column=col_idx, value=col_name)
            
            # Write data
            for row_idx, row in enumerate(df.itertuples(index=False), 2):
                for col_idx, value in enumerate(row, 1):
                    sheet.cell(row=row_idx, column=col_idx, value=str(value))

        except Exception:
            # If a table fails to process, skip it
            continue

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return output


def extract_tables_to_tsv(image_bytes: bytes) -> str:
    """
    Extract detected tables from an image and return TSV text.
    """
    ocr = TesseractOCR(n_threads=1, lang="eng")
    image = Image(src=image_bytes)

    tables = image.extract_tables(
        ocr=ocr,
        implicit_rows=True,
        min_confidence=50,
    )

    if not tables:
        return "No tables detected in the image"

    sections = []
    for idx, table in enumerate(tables, start=1):
        try:
            df = table.df
            # Convert all data to string to be safe
            df = df.astype(str)
            
            # Convert to TSV, including header
            table_text = df.to_csv(sep='\t', index=False, header=True)

            if table_text.strip():
                sections.append(f"Table {idx}\n{table_text}")
        except Exception:
            # If a table fails to process, skip it
            continue

    if not sections:
        return "Could not extract any table content."

    return "\n\n".join(sections)
